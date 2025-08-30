from typing import List, Dict, Any, Optional
from pathlib import Path
import json
import zipfile
import tempfile
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from .models import SessionPackage, SessionMetadata, PackageRequest
from .sop_generator import SOPGenerator
from .irb_generator import IRBGenerator
from .bids_generator import BIDSGenerator
from .common_utils.database import DatabaseConnection
from .common_utils.logging_config import get_logger

logger = get_logger(__name__)


class SessionPackager:
    """Main session packaging engine that combines all components."""

    def __init__(self, storage_path: Optional[Path] = None):
        """Initialize the session packager with all generators."""
        self.sop_generator = SOPGenerator()
        self.irb_generator = IRBGenerator()
        self.bids_generator = BIDSGenerator()

        # Set up storage
        if storage_path is None:
            storage_path = Path("./data/packages.db")
        self.storage_db = DatabaseConnection(storage_path)
        self._init_storage()

    def _init_storage(self):
        """Initialize the storage database with required tables."""
        schema = {
            "packages": {
                "id": "INTEGER PRIMARY KEY AUTOINCREMENT",
                "session_id": "TEXT UNIQUE NOT NULL",
                "package_data": "TEXT NOT NULL",  # JSON string
                "created_at": "TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
                "updated_at": "TIMESTAMP DEFAULT CURRENT_TIMESTAMP"
            }
        }

        # Create tables if they don't exist
        for table_name, table_schema in schema.items():
            self.storage_db.create_table(table_name, table_schema)

    def save_package(self, package: SessionPackage) -> bool:
        """Save a session package to storage."""
        try:
            # Convert package to dict for storage
            package_dict = {
                "session_metadata": package.session_metadata.model_dump(),
                "bids_events": [event.model_dump() for event in package.bids_events],
                "sop_documents": [doc.model_dump() for doc in package.sop_documents],
                "irb_documents": [doc.model_dump() for doc in package.irb_documents],
                "created_at": package.created_at.isoformat()
            }

            package_json = json.dumps(package_dict, default=str)

            # Check if package already exists
            existing = self.storage_db.execute_query(
                "SELECT id FROM packages WHERE session_id = ?",
                (package.session_metadata.session_id,)
            )

            if existing:
                # Update existing
                self.storage_db.execute_update(
                    "UPDATE packages SET package_data = ?, updated_at = CURRENT_TIMESTAMP WHERE session_id = ?",
                    (package_json, package.session_metadata.session_id)
                )
            else:
                # Insert new
                self.storage_db.execute_update(
                    "INSERT INTO packages (session_id, package_data) VALUES (?, ?)",
                    (package.session_metadata.session_id, package_json)
                )

            logger.info(f"Package saved: {package.session_metadata.session_id}")
            return True

        except Exception as e:
            logger.error(f"Error saving package: {e}")
            return False

    def load_package(self, session_id: str) -> Optional[SessionPackage]:
        """Load a session package from storage."""
        try:
            results = self.storage_db.execute_query(
                "SELECT package_data FROM packages WHERE session_id = ?",
                (session_id,)
            )

            if not results:
                return None

            package_dict = json.loads(results[0]["package_data"])

            # Reconstruct objects
            session_metadata = SessionMetadata(**package_dict["session_metadata"])

            return SessionPackage(
                session_metadata=session_metadata,
                bids_events=package_dict.get("bids_events", []),
                sop_documents=package_dict.get("sop_documents", []),
                irb_documents=package_dict.get("irb_documents", []),
                created_at=datetime.fromisoformat(package_dict["created_at"])
            )

        except Exception as e:
            logger.error(f"Error loading package {session_id}: {e}")
            return None

    def list_packages(self) -> List[Dict[str, Any]]:
        """List all saved packages with basic info."""
        try:
            results = self.storage_db.execute_query(
                "SELECT session_id, created_at FROM packages ORDER BY created_at DESC"
            )
            return results
        except Exception as e:
            logger.error(f"Error listing packages: {e}")
            return []

    def delete_package(self, session_id: str) -> bool:
        """Delete a package from storage."""
        try:
            # Check if package exists first
            existing = self.storage_db.execute_query(
                "SELECT id FROM packages WHERE session_id = ?",
                (session_id,)
            )
            
            if not existing:
                return False
            
            result = self.storage_db.execute_update(
                "DELETE FROM packages WHERE session_id = ?",
                (session_id,)
            )
            
            logger.info(f"Package deleted: {session_id}")
            return result > 0
        except Exception as e:
            logger.error(f"Error deleting package {session_id}: {e}")
            return False
    
    def create_session_package(self, request: PackageRequest, save_to_storage: bool = True) -> SessionPackage:
        """Create a complete session package based on the request."""

        # Generate SOP documents
        sop_documents = []
        if request.include_sop:
            sop_documents.append(self.sop_generator.generate_sop(request.session_metadata))

        # Generate IRB documents
        irb_documents = []
        if request.include_irb:
            irb_documents = self.irb_generator.generate_irb_package(request.session_metadata)

        # Generate BIDS events
        bids_events = []
        if request.include_bids:
            bids_events = self.bids_generator._generate_events_template(
                request.session_metadata,
                request.custom_events
            )

        package = SessionPackage(
            session_metadata=request.session_metadata,
            bids_events=bids_events,
            sop_documents=sop_documents,
            irb_documents=irb_documents,
            created_at=datetime.now()
        )

        # Save to storage if requested
        if save_to_storage:
            self.save_package(package)

        return package
    
    def export_package(self, package: SessionPackage, 
                      output_dir: Path, 
                      formats: List[str] = None) -> Dict[str, Path]:
        """Export the package in various formats."""
        
        if formats is None:
            formats = ["json", "pdf", "docx", "bids", "zip"]
        
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        exported_files = {}
        
        # Export JSON summary
        if "json" in formats:
            json_path = self._export_json(package, output_dir)
            exported_files["json"] = json_path
        
        # Export PDF documents
        if "pdf" in formats:
            pdf_path = self._export_pdf(package, output_dir)
            exported_files["pdf"] = pdf_path
        
        # Export Word documents
        if "docx" in formats:
            docx_path = self._export_docx(package, output_dir)
            exported_files["docx"] = docx_path
        
        # Export BIDS structure
        if "bids" in formats:
            bids_path = self._export_bids(package, output_dir)
            exported_files["bids"] = bids_path
        
        # Export ZIP package
        if "zip" in formats:
            zip_path = self._export_zip(package, output_dir)
            exported_files["zip"] = zip_path
        
        return exported_files
    
    def _export_json(self, package: SessionPackage, output_dir: Path) -> Path:
        """Export package as JSON."""
        json_path = output_dir / f"{package.session_metadata.session_id}_package.json"
        
        # Convert to dict for JSON serialization
        package_dict = {
            "session_metadata": package.session_metadata.model_dump(),
            "bids_events": [event.model_dump() for event in package.bids_events],
            "sop_documents": [doc.model_dump() for doc in package.sop_documents],
            "irb_documents": [doc.model_dump() for doc in package.irb_documents],
            "created_at": package.created_at.isoformat()
        }
        
        json_path.write_text(json.dumps(package_dict, indent=2, default=str))
        return json_path
    
    def _export_pdf(self, package: SessionPackage, output_dir: Path) -> Path:
        """Export package as PDF documents."""
        pdf_path = output_dir / f"{package.session_metadata.session_id}_documents.pdf"
        
        doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title page
        title = f"Session Package: {package.session_metadata.study_name}"
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))
        
        subtitle = f"Session ID: {package.session_metadata.session_id}"
        story.append(Paragraph(subtitle, styles['Heading2']))
        story.append(Spacer(1, 12))
        
        # Session metadata
        story.append(Paragraph("Session Information", styles['Heading1']))
        metadata_text = f"""
        <b>Study Name:</b> {package.session_metadata.study_name}<br/>
        <b>Principal Investigator:</b> {package.session_metadata.principal_investigator}<br/>
        <b>Modality:</b> {package.session_metadata.modality.value}<br/>
        <b>Session Type:</b> {package.session_metadata.session_type.value}<br/>
        <b>Population:</b> {package.session_metadata.participant_population.value}<br/>
        <b>Duration:</b> {package.session_metadata.duration_minutes} minutes<br/>
        <b>Expected Participants:</b> {package.session_metadata.expected_participants}<br/>
        <b>Risk Level:</b> {package.session_metadata.risk_level.value}<br/>
        """
        story.append(Paragraph(metadata_text, styles['Normal']))
        story.append(PageBreak())
        
        # SOP Documents
        if package.sop_documents:
            story.append(Paragraph("Standard Operating Procedures", styles['Heading1']))
            for sop in package.sop_documents:
                story.append(Paragraph(sop.title, styles['Heading2']))
                story.append(Paragraph(f"<b>Purpose:</b> {sop.purpose}", styles['Normal']))
                story.append(Paragraph(f"<b>Scope:</b> {sop.scope}", styles['Normal']))
                
                story.append(Paragraph("<b>Procedure Steps:</b>", styles['Heading3']))
                for step in sop.procedure_steps:
                    story.append(Paragraph(f"• {step}", styles['Normal']))
                
                story.append(Paragraph("<b>Safety Considerations:</b>", styles['Heading3']))
                for safety in sop.safety_considerations:
                    story.append(Paragraph(f"• {safety}", styles['Normal']))
                
                story.append(PageBreak())
        
        # IRB Documents
        if package.irb_documents:
            story.append(Paragraph("IRB Documents", styles['Heading1']))
            for doc in package.irb_documents:
                story.append(Paragraph(doc.document_type.replace('_', ' ').title(), styles['Heading2']))
                story.append(Paragraph(doc.content, styles['Normal']))
                story.append(PageBreak())
        
        doc.build(story)
        return pdf_path
    
    def _export_docx(self, package: SessionPackage, output_dir: Path) -> Path:
        """Export package as Word documents."""
        docx_path = output_dir / f"{package.session_metadata.session_id}_documents.docx"
        
        doc = Document()
        doc.add_heading(f"Session Package: {package.session_metadata.study_name}", 0)
        
        # Session metadata
        doc.add_heading("Session Information", 1)
        doc.add_paragraph(f"Study Name: {package.session_metadata.study_name}")
        doc.add_paragraph(f"Principal Investigator: {package.session_metadata.principal_investigator}")
        doc.add_paragraph(f"Modality: {package.session_metadata.modality.value}")
        doc.add_paragraph(f"Session Type: {package.session_metadata.session_type.value}")
        doc.add_paragraph(f"Population: {package.session_metadata.participant_population.value}")
        doc.add_paragraph(f"Duration: {package.session_metadata.duration_minutes} minutes")
        doc.add_paragraph(f"Expected Participants: {package.session_metadata.expected_participants}")
        doc.add_paragraph(f"Risk Level: {package.session_metadata.risk_level.value}")
        doc.add_page_break()
        
        # SOP Documents
        if package.sop_documents:
            doc.add_heading("Standard Operating Procedures", 1)
            for sop in package.sop_documents:
                doc.add_heading(sop.title, 2)
                doc.add_paragraph(f"Purpose: {sop.purpose}")
                doc.add_paragraph(f"Scope: {sop.scope}")
                
                doc.add_heading("Procedure Steps", 3)
                for step in sop.procedure_steps:
                    doc.add_paragraph(f"• {step}", style='List Bullet')
                
                doc.add_heading("Safety Considerations", 3)
                for safety in sop.safety_considerations:
                    doc.add_paragraph(f"• {safety}", style='List Bullet')
                
                doc.add_page_break()
        
        # IRB Documents
        if package.irb_documents:
            doc.add_heading("IRB Documents", 1)
            for doc in package.irb_documents:
                doc.add_heading(doc.document_type.replace('_', ' ').title(), 2)
                doc.add_paragraph(doc.content)
                doc.add_page_break()
        
        doc.save(str(docx_path))
        return docx_path
    
    def _export_bids(self, package: SessionPackage, output_dir: Path) -> Path:
        """Export BIDS structure."""
        bids_dir = output_dir / f"{package.session_metadata.session_id}_bids"
        self.bids_generator.export_bids_package(
            package.session_metadata, 
            bids_dir, 
            package.bids_events
        )
        return bids_dir
    
    def _export_zip(self, package: SessionPackage, output_dir: Path) -> Path:
        """Export complete package as ZIP file."""
        zip_path = output_dir / f"{package.session_metadata.session_id}_complete_package.zip"
        
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Export all formats to temp directory
            self.export_package(package, temp_path, ["json", "pdf", "docx", "bids"])
            
            # Create ZIP
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, dirs, files in temp_path.rglob('*'):
                    for file in files:
                        file_path = root / file
                        arcname = file_path.relative_to(temp_path)
                        zipf.write(file_path, arcname)
        
        return zip_path
    
    def validate_package(self, package: SessionPackage) -> Dict[str, Any]:
        """Validate the session package for completeness and compliance."""
        
        validation_results = {
            "valid": True,
            "warnings": [],
            "errors": []
        }
        
        # Validate session metadata
        if not package.session_metadata.session_id:
            validation_results["errors"].append("Session ID is required")
            validation_results["valid"] = False
        
        if not package.session_metadata.study_name:
            validation_results["errors"].append("Study name is required")
            validation_results["valid"] = False
        
        if not package.session_metadata.principal_investigator:
            validation_results["errors"].append("Principal investigator is required")
            validation_results["valid"] = False
        
        # Validate BIDS events
        for event in package.bids_events:
            if event.onset < 0:
                validation_results["errors"].append("Event onset cannot be negative")
                validation_results["valid"] = False
            
            if event.duration < 0:
                validation_results["errors"].append("Event duration cannot be negative")
                validation_results["valid"] = False
        
        # Validate SOP documents
        for sop in package.sop_documents:
            if not sop.title:
                validation_results["warnings"].append("SOP document missing title")
            
            if not sop.procedure_steps:
                validation_results["warnings"].append("SOP document missing procedure steps")
        
        # Validate IRB documents
        for doc in package.irb_documents:
            if not doc.document_type:
                validation_results["warnings"].append("IRB document missing type")
            
            if not doc.content:
                validation_results["warnings"].append("IRB document missing content")
        
        return validation_results
    
    def get_package_summary(self, package: SessionPackage) -> Dict[str, Any]:
        """Get a summary of the package contents."""
        
        return {
            "session_id": package.session_metadata.session_id,
            "study_name": package.session_metadata.study_name,
            "modality": package.session_metadata.modality.value,
            "session_type": package.session_metadata.session_type.value,
            "created_at": package.created_at.isoformat(),
            "document_counts": {
                "sop_documents": len(package.sop_documents),
                "irb_documents": len(package.irb_documents),
                "bids_events": len(package.bids_events)
            },
            "total_duration": package.session_metadata.duration_minutes,
            "expected_participants": package.session_metadata.expected_participants,
            "risk_level": package.session_metadata.risk_level.value
        }
